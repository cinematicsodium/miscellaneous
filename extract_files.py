import os
import re
import pandas as pd
from docx import Document

def extract_info(file_path):
    doc = Document(file_path)
    text = ' '.join([paragraph.text for paragraph in doc.paragraphs])
    
    name = re.search('Name: (.*?) ', text).group(1)
    dob = re.search('Date of Birth: (.*?) ', text).group(1)
    employment_history = re.search('Employment History: (.*?) ', text).group(1)
    
    return {'Name': name, 'Date of Birth': dob, 'Employment History': employment_history}

df = pd.DataFrame(columns=['Name', 'Date of Birth', 'Employment History'])

for file in os.listdir('/path/to/your/directory'):
    if file.endswith('.rtf'):
        info = extract_info(os.path.join('/path/to/your/directory', file))
        df = df.append(info, ignore_index=True)

df.to_excel('output.xlsx', index=False)


def extract_info(file_path):
    doc = Document(file_path)
    text = ' '.join([paragraph.text for paragraph in doc.paragraphs])

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += ' ' + cell.text

    name = re.search('Name: (.*?) ', text).group(1)
    dob = re.search('Date of Birth: (.*?) ', text).group(1)
    employment_history = re.search('Employment History: (.*?) ', text).group(1)
    
    return {'Name': name, 'Date of Birth': dob, 'Employment History': employment_history}


def extract_info(file_path):
    doc = Document(file_path)
    text = ' '.join([paragraph.text for paragraph in doc.paragraphs])

    # Also extract text from tables
    info = {}
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            for i in range(0, len(cells), 2):  # assuming each key-value pair is in two cells
                key = cells[i].text.strip()
                value = cells[i+1].text.strip() if i+1 < len(cells) else ''
                info[key] = value

    return info